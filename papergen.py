"""
PaperGen - Academic Paper Template Generator
Supports .docx, .typ (Typst), and .tex (LaTeX) templates.
"""

import sys
import os
import json
from pathlib import Path
from datetime import date
import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document

APPDATA = Path(os.environ.get("APPDATA", Path.home())) / "PaperGen"
CLASSES_FILE = APPDATA / "classes.json"
PROFILE_FILE = APPDATA / "profile.json"
CONFIG_FILE = APPDATA / "config.json"
TEMPLATES_DIR = APPDATA / "templates"

APPDATA.mkdir(parents=True, exist_ok=True)
TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

SUPPORTED_EXTS = (".docx", ".tex", ".typ")
DEFAULT_FILENAME_FORMAT = "{{last name}}{{title}}"


# ── Profile ───────────────────────────────────────────────────────────────────
def load_profile():
    if not PROFILE_FILE.exists():
        return {"name_first": "", "name_last": "", "university": ""}
    with open(PROFILE_FILE, "r", encoding="utf-8") as f:
        return json.load(f)


def save_profile(profile):
    with open(PROFILE_FILE, "w", encoding="utf-8") as f:
        json.dump(profile, f, indent=2, ensure_ascii=False)


def load_config():
    if not CONFIG_FILE.exists():
        return {"filename_format": DEFAULT_FILENAME_FORMAT}
    with open(CONFIG_FILE, "r", encoding="utf-8") as f:
        config = json.load(f)
    filename_format = config.get("filename_format", DEFAULT_FILENAME_FORMAT)
    if not isinstance(filename_format, str) or not filename_format.strip():
        filename_format = DEFAULT_FILENAME_FORMAT
    return {"filename_format": filename_format}


def save_config(config):
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(config, f, indent=2, ensure_ascii=False)


# ── Classes ───────────────────────────────────────────────────────────────────
def migrate_class(cls):
    changed = False
    if "class_code" not in cls:
        old = cls.get("class_number", "")
        parts = old.split(" ", 1)
        cls["class_code"] = parts[0] if parts else ""
        cls["class_number"] = parts[1] if len(parts) > 1 else ""
        changed = True
    if "university" not in cls:
        cls["university"] = ""
        changed = True
    return cls, changed


def load_classes():
    if not CLASSES_FILE.exists():
        return []
    with open(CLASSES_FILE, "r", encoding="utf-8") as f:
        classes = json.load(f)
    migrated = [migrate_class(c) for c in classes]
    classes_out = [c for c, _ in migrated]
    if any(changed for _, changed in migrated):
        save_classes(classes_out)
    return classes_out


def save_classes(classes):
    with open(CLASSES_FILE, "w", encoding="utf-8") as f:
        json.dump(classes, f, indent=2, ensure_ascii=False)


def list_templates():
    """Return templates grouped by type, alphabetized within each type."""
    templates = []
    for ext in SUPPORTED_EXTS:
        names = (p.name for p in TEMPLATES_DIR.glob(f"*{ext}"))
        templates.extend(sorted(names, key=str.casefold))
    return templates


def class_label(cls):
    code   = cls.get("class_code", "")
    num    = cls.get("class_number", "")
    name   = cls.get("class_name", "")
    prefix = f"{code} {num}".strip()
    return f"{prefix} \u2014 {name}" if prefix else name


# ── Document generation ───────────────────────────────────────────────────────
def build_placeholders(cls, title, due_date):
    profile    = load_profile()
    class_full = f"{cls.get('class_code', '')} {cls.get('class_number', '')}".strip()
    return {
        "{{your_name}}":      cls.get("your_name", ""),
        "{{name_first}}":     profile.get("name_first", ""),
        "{{name_last}}":      profile.get("name_last", ""),
        "{{professor_name}}": cls.get("professor_name", ""),
        "{{class_name}}":     cls.get("class_name", ""),
        "{{class_code}}":     cls.get("class_code", ""),
        "{{class_number}}":   cls.get("class_number", ""),
        "{{class_full}}":     class_full,
        "{{school}}":         cls.get("school", ""),
        "{{university}}":     cls.get("university", "") or profile.get("university", ""),
        "{{paper_title}}":    title,
        "{{due_date}}":       due_date,
    }


def build_output_name(cls, title, ext):
    profile = load_profile()
    first   = profile.get("name_first", "").strip()
    last    = profile.get("name_last", "").strip()
    if not first or not last:
        parts = cls.get("your_name", "").strip().split()
        last  = parts[-1] if len(parts) >= 2 else "".join(parts)
        first = parts[0]  if len(parts) >= 2 else ""
    safe_title  = "".join(c for c in title if c.isalnum() or c in " ").title().replace(" ", "")
    filename_format = load_config()["filename_format"]
    replacements = {
        "{{last name}}":  last,
        "{{first name}}": first,
        "{{class code}}":  cls.get("class_code", ""),
        "{{class number}}": cls.get("class_number", ""),
        "{{title}}":      safe_title,
    }
    filename = filename_format
    for placeholder, value in replacements.items():
        filename = filename.replace(placeholder, value)

    invalid = '<>:"/\\|?*'
    filename = "".join(c for c in filename if c not in invalid and ord(c) >= 32)
    filename = filename.strip().rstrip(".")
    if not filename:
        filename = safe_title or "paper"
    return f"{filename}{ext}"


def replace_in_paragraph(para, placeholders):
    """Handles Word splitting placeholders across multiple runs."""
    if not para.runs:
        return
    full_text = "".join(run.text for run in para.runs)
    replaced  = full_text
    for key, val in placeholders.items():
        replaced = replaced.replace(key, val)
    if replaced == full_text:
        return
    para.runs[0].text = replaced
    for run in para.runs[1:]:
        run.text = ""


def generate_docx(cls, template_path, placeholders, output_path):
    doc = Document(str(template_path))
    for para in doc.paragraphs:
        replace_in_paragraph(para, placeholders)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, placeholders)
    for section in doc.sections:
        for para in section.header.paragraphs:
            replace_in_paragraph(para, placeholders)
        for para in section.footer.paragraphs:
            replace_in_paragraph(para, placeholders)
    doc.save(str(output_path))


def generate_text(template_path, placeholders, output_path):
    """Plain-text substitution for .typ and .tex templates."""
    text = template_path.read_text(encoding="utf-8")
    for key, val in placeholders.items():
        text = text.replace(key, val)
    output_path.write_text(text, encoding="utf-8")


def generate_paper(cls, template_filename, title, due_date, output_folder):
    """Dispatch to the right generator based on template extension."""
    template_path = TEMPLATES_DIR / template_filename
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    ext          = template_path.suffix.lower()
    placeholders = build_placeholders(cls, title, due_date)
    out_path     = Path(output_folder) / build_output_name(cls, title, ext)

    if ext == ".docx":
        generate_docx(cls, template_path, placeholders, out_path)
    elif ext in (".typ", ".tex"):
        generate_text(template_path, placeholders, out_path)
    else:
        raise ValueError(f"Unsupported template type: {ext}")

    return out_path


# ── Minimal UI helpers ────────────────────────────────────────────────────────
def lbl(parent, text, bold=False, **kw):
    return tk.Label(parent, text=text,
                    font=("Segoe UI", 10, "bold" if bold else "normal"), **kw)

def entry(parent, textvariable=None, width=32):
    return tk.Entry(parent, textvariable=textvariable, width=width,
                    font=("Segoe UI", 10))

def dropdown(parent, values, textvariable=None, width=30):
    return ttk.Combobox(parent, values=values, textvariable=textvariable,
                        width=width, state="readonly", font=("Segoe UI", 10))

def btn(parent, text, command, **kw):
    return tk.Button(parent, text=text, command=command,
                     font=("Segoe UI", 10), padx=10, pady=4, **kw)

def hsep(parent):
    return ttk.Separator(parent, orient="horizontal")


# ── New Paper Window ──────────────────────────────────────────────────────────
class NewPaperWindow:
    def __init__(self, output_folder):
        self.output_folder = output_folder
        self.root = tk.Tk()
        self.root.title("PaperGen \u2014 New Paper")
        self.root.resizable(False, False)
        self._build()
        self._center()

    def _build(self):
        root = self.root

        lbl(root, "New Paper", bold=True).grid(
            row=0, column=0, columnspan=2, pady=(12, 0), padx=16, sticky="w")
        lbl(root, f"\u2192 {self.output_folder}", fg="gray").grid(
            row=1, column=0, columnspan=2, padx=16, sticky="w")
        hsep(root).grid(row=2, column=0, columnspan=2, sticky="ew", padx=16, pady=6)

        classes   = load_classes()
        templates = list_templates()

        self.class_var    = tk.StringVar()
        self.template_var = tk.StringVar()
        self.title_var    = tk.StringVar()
        self.date_var     = tk.StringVar(value=date.today().strftime("%B %d, %Y"))

        if not classes:
            messagebox.showwarning("No Classes",
                "No classes found. Add one via Configuration.", parent=root)
        if not templates:
            messagebox.showwarning("No Templates",
                f"No templates found in:\n{TEMPLATES_DIR}\n\n"
                "Supported: .docx, .typ, .tex", parent=root)

        class_labels = [class_label(c) for c in classes]
        self._classes = classes

        self._class_dropdown = dropdown(root, class_labels, self.class_var)
        fields = [
            ("Class",            self._class_dropdown),
            ("Template / Style", dropdown(root, templates, self.template_var)),
            ("Paper Title",      entry(root, self.title_var)),
            ("Due Date",         entry(root, self.date_var)),
        ]

        for i, (label_text, widget) in enumerate(fields):
            lbl(root, label_text).grid(
                row=3+i*2, column=0, sticky="w", padx=16, pady=(8, 1))
            widget.grid(
                row=4+i*2, column=0, columnspan=2, sticky="ew", padx=16, pady=(0, 4))

        root.grid_columnconfigure(0, weight=1)

        if class_labels:
            self.class_var.set(class_labels[0])
        if templates:
            self.template_var.set(templates[0])

        r = 3 + len(fields) * 2
        hsep(root).grid(row=r, column=0, columnspan=2, sticky="ew", padx=16, pady=8)

        btns = tk.Frame(root)
        btns.grid(row=r+1, column=0, columnspan=2, padx=16, pady=(0, 12), sticky="ew")
        btn(btns, "Configuration", self._open_configuration).pack(side="left")
        btn(btns, "Generate", self._generate).pack(side="right")

    def _center(self):
        self.root.update_idletasks()
        w, h = self.root.winfo_width(), self.root.winfo_height()
        sw, sh = self.root.winfo_screenwidth(), self.root.winfo_screenheight()
        self.root.geometry(f"{w}x{h}+{(sw-w)//2}+{(sh-h)//2}")

    def _open_configuration(self):
        self.root.withdraw()
        ConfigurationWindow(on_close=self._on_configuration_close)

    def _on_configuration_close(self):
        self.root.deiconify()
        classes = load_classes()
        class_labels = [class_label(c) for c in classes]
        self._classes = classes
        self._class_dropdown["values"] = class_labels
        if class_labels:
            self.class_var.set(class_labels[0])
        else:
            self.class_var.set("")

    def _generate(self):
        classes      = load_classes()
        class_labels = [class_label(c) for c in classes]

        if not self.class_var.get():
            messagebox.showerror("Error", "Select a class.", parent=self.root); return
        if not self.template_var.get():
            messagebox.showerror("Error", "Select a template.", parent=self.root); return
        if not self.title_var.get().strip():
            messagebox.showerror("Error", "Enter a paper title.", parent=self.root); return

        idx = class_labels.index(self.class_var.get())
        try:
            out = generate_paper(
                cls=classes[idx],
                template_filename=self.template_var.get(),
                title=self.title_var.get().strip(),
                due_date=self.date_var.get().strip(),
                output_folder=self.output_folder,
            )
            messagebox.showinfo("Done", f"Created:\n{out.name}", parent=self.root)
            self.root.destroy()
        except Exception as e:
            messagebox.showerror("Error", str(e), parent=self.root)

    def run(self):
        self.root.mainloop()


# ── Configuration Window ─────────────────────────────────────────────────────
class ConfigurationWindow:
    def __init__(self, on_close=None):
        self.on_close = on_close
        self.root = tk.Toplevel() if on_close else tk.Tk()
        self.root.title("PaperGen \u2014 Configuration")
        self.root.resizable(False, False)
        self._build()
        self._center()
        self.root.protocol("WM_DELETE_WINDOW", self._on_close)

    def _build(self):
        root = self.root
        lbl(root, "Configuration", bold=True).pack(padx=16, pady=(12, 2), anchor="w")
        lbl(root, f"Classes: {CLASSES_FILE}", fg="gray").pack(padx=16, anchor="w")
        lbl(root, f"Settings: {CONFIG_FILE}", fg="gray").pack(padx=16, anchor="w")
        hsep(root).pack(fill="x", padx=16, pady=6)

        lbl(root, "Classes", bold=True).pack(padx=16, anchor="w")
        self.list_frame = tk.Frame(root)
        self.list_frame.pack(padx=16, fill="both")
        self._render_list()

        class_btns = tk.Frame(root)
        class_btns.pack(padx=16, pady=(6, 0), fill="x")
        btn(class_btns, "+ Add Class",    self._add).pack(side="left")
        btn(class_btns, "Edit Profile",   self._edit_profile).pack(side="left", padx=(8, 0))
        btn(class_btns, "Clear Semester", self._clear, fg="red").pack(side="right")

        hsep(root).pack(fill="x", padx=16, pady=10)
        lbl(root, "Filename Convention", bold=True).pack(padx=16, anchor="w")
        lbl(root, "Tokens: {{last name}}, {{first name}}, {{class code}}, {{class number}}, {{title}}",
            fg="gray").pack(padx=16, anchor="w")

        self.filename_format_var = tk.StringVar(
            value=load_config()["filename_format"])
        filename_form = tk.Frame(root)
        filename_form.pack(padx=16, pady=(8, 0), fill="x")
        lbl(filename_form, "Filename format").pack(side="left")
        entry(filename_form, self.filename_format_var, width=34).pack(
            side="left", padx=(10, 0), fill="x", expand=True)

        config_btns = tk.Frame(root)
        config_btns.pack(padx=16, pady=(6, 12), fill="x")
        btn(config_btns, "Save Filename Format", self._save_filename_format).pack(side="right")

    def _save_filename_format(self):
        filename_format = self.filename_format_var.get().strip()
        if not filename_format:
            messagebox.showerror(
                "Required", "Filename format cannot be empty.", parent=self.root)
            return
        save_config({"filename_format": filename_format})
        messagebox.showinfo("Saved", "Filename convention saved.", parent=self.root)

    def _render_list(self):
        for w in self.list_frame.winfo_children():
            w.destroy()
        classes = load_classes()
        if not classes:
            lbl(self.list_frame, "No classes yet.").pack(pady=8)
            return
        for i, cls in enumerate(classes):
            row = tk.Frame(self.list_frame)
            row.pack(fill="x", pady=2)
            info = tk.Frame(row)
            info.pack(side="left")
            lbl(info, class_label(cls), bold=True).pack(anchor="w")
            sub = f"{cls.get('professor_name', '')} \u00b7 {cls.get('school', '')}"
            if cls.get("university"):
                sub += f" \u00b7 {cls['university']}"
            lbl(info, sub, fg="gray").pack(anchor="w")
            acts = tk.Frame(row)
            acts.pack(side="right")
            btn(acts, "Delete", lambda i=i: self._delete(i), fg="red").pack(side="right", padx=2)
            btn(acts, "Edit",   lambda i=i: self._edit(i)).pack(side="right", padx=2)

    def _add(self):
        ClassFormDialog(self.root, on_save=self._on_save)

    def _edit(self, idx):
        ClassFormDialog(self.root, on_save=self._on_save, existing=load_classes()[idx], idx=idx)

    def _edit_profile(self):
        ProfileFormDialog(self.root)

    def _on_save(self, data, idx=None):
        classes = load_classes()
        if idx is None:
            classes.append(data)
        else:
            classes[idx] = data
        save_classes(classes)
        self._render_list()

    def _delete(self, idx):
        classes = load_classes()
        if messagebox.askyesno("Delete", f"Delete {class_label(classes[idx])}?", parent=self.root):
            classes.pop(idx)
            save_classes(classes)
            self._render_list()

    def _clear(self):
        classes = load_classes()
        if not classes:
            messagebox.showinfo("Empty", "No classes to clear.", parent=self.root); return
        if messagebox.askyesno("Clear Semester", f"Delete all {len(classes)} class(es)?",
                               icon="warning", parent=self.root):
            save_classes([])
            self._render_list()

    def _center(self):
        self.root.update_idletasks()
        w, h = self.root.winfo_width(), self.root.winfo_height()
        sw, sh = self.root.winfo_screenwidth(), self.root.winfo_screenheight()
        self.root.geometry(f"{w}x{h}+{(sw-w)//2}+{(sh-h)//2}")

    def _on_close(self):
        if self.on_close:
            self.on_close()
        self.root.destroy()

    def run(self):
        self.root.mainloop()


# ── Profile Form Dialog ───────────────────────────────────────────────────────
class ProfileFormDialog(tk.Toplevel):
    FIELDS = [
        ("name_first", "First Name"),
        ("name_last",  "Last Name"),
        ("university", "University (external)"),
    ]

    def __init__(self, parent):
        super().__init__(parent)
        self.title("Edit Profile")
        self.resizable(True, False)
        self.grab_set()

        profile = load_profile()
        self.vars = {k: tk.StringVar(value=profile.get(k, "")) for k, _ in self.FIELDS}

        lbl(self, "Profile", bold=True).pack(padx=16, pady=(12, 2), anchor="w")
        lbl(self, "Pre-fills new classes. Provides {{name_first}} and {{name_last}}.",
            fg="gray").pack(padx=16, anchor="w")

        form = tk.Frame(self)
        form.pack(padx=16, pady=(8, 0), fill="x", expand=True)
        form.grid_columnconfigure(0, weight=1)
        for i, (key, display) in enumerate(self.FIELDS):
            lbl(form, display).grid(row=i*2, column=0, sticky="w", pady=(6, 1))
            entry(form, self.vars[key], width=48).grid(row=i*2+1, column=0, sticky="ew")

        hsep(self).pack(fill="x", padx=16, pady=10)
        btns = tk.Frame(self)
        btns.pack(padx=16, pady=(0, 12), fill="x")
        btn(btns, "Cancel", self.destroy).pack(side="left")
        btn(btns, "Save", self._save).pack(side="right")

        self.update_idletasks()
        w, h = self.winfo_width(), self.winfo_height()
        sw, sh = self.winfo_screenwidth(), self.winfo_screenheight()
        self.geometry(f"{w}x{h}+{(sw-w)//2}+{(sh-h)//2}")

    def _save(self):
        save_profile({k: v.get().strip() for k, v in self.vars.items()})
        self.destroy()


# ── Class Form Dialog ─────────────────────────────────────────────────────────
class ClassFormDialog(tk.Toplevel):
    FIELDS = [
        ("your_name",      "Your Name"),
        ("professor_name", "Professor Name"),
        ("class_name",     "Class Name"),
        ("class_code",     "Class Code (e.g. PHIL)"),
        ("class_number",   "Class Number (e.g. 101)"),
        ("school",         "School / College (internal)"),
        ("university",     "University (external)"),
    ]

    def __init__(self, parent, on_save, existing=None, idx=None):
        super().__init__(parent)
        self.on_save = on_save
        self.idx = idx
        self.title("Edit Class" if existing else "Add Class")
        self.resizable(True, False)
        self.grab_set()

        profile = load_profile()
        profile_name = f"{profile.get('name_first', '')} {profile.get('name_last', '')}".strip()
        defaults = {
            "your_name":  profile_name,
            "university": profile.get("university", ""),
        }

        self.vars = {
            k: tk.StringVar(value=existing.get(k, "") if existing else defaults.get(k, ""))
            for k, _ in self.FIELDS
        }

        lbl(self, "Edit Class" if existing else "Add Class", bold=True).pack(
            padx=16, pady=(12, 8), anchor="w")

        form = tk.Frame(self)
        form.pack(padx=16, fill="x", expand=True)
        form.grid_columnconfigure(0, weight=1)
        for i, (key, display) in enumerate(self.FIELDS):
            lbl(form, display).grid(row=i*2, column=0, sticky="w", pady=(6, 1))
            entry(form, self.vars[key], width=48).grid(row=i*2+1, column=0, sticky="ew")

        hsep(self).pack(fill="x", padx=16, pady=10)
        btns = tk.Frame(self)
        btns.pack(padx=16, pady=(0, 12), fill="x")
        btn(btns, "Cancel", self.destroy).pack(side="left")
        btn(btns, "Save", self._save).pack(side="right")

        self.update_idletasks()
        w, h = self.winfo_width(), self.winfo_height()
        sw, sh = self.winfo_screenwidth(), self.winfo_screenheight()
        self.geometry(f"{w}x{h}+{(sw-w)//2}+{(sh-h)//2}")

    def _save(self):
        data = {k: v.get().strip() for k, v in self.vars.items()}
        if not data["class_name"]:
            messagebox.showerror("Required", "Class name is required.", parent=self)
            return
        self.on_save(data, self.idx)
        self.destroy()


# ── Entry point ───────────────────────────────────────────────────────────────
if __name__ == "__main__":
    if len(sys.argv) >= 2 and sys.argv[1] not in ("--manage", "--config"):
        output_folder = sys.argv[1]
        if not os.path.isdir(output_folder):
            messagebox.showerror("Error", f"Invalid folder:\n{output_folder}")
            sys.exit(1)
        NewPaperWindow(output_folder).run()
    else:
        ConfigurationWindow().run()
